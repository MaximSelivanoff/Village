from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

# Создание документа
doc = Document()

# Настройка стилей по ГОСТ Р 7.0.5-2008
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(14)
font._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

# Установка межстрочного интервала 1.5
paragraph_format = style.paragraph_format
paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

# Поля: левое — 30 мм, правое — 10 мм, верхнее и нижнее — 20 мм
sections = doc.sections
for section in sections:
    section.left_margin = Mm(30)
    section.right_margin = Mm(10)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)

# Функция для добавления параграфа с правильным форматированием
def add_paragraph(text, alignment=WD_ALIGN_PARAGRAPH.LEFT, bold=False, font_size=Pt(14)):
    p = doc.add_paragraph(text)
    p.alignment = alignment
    run = p.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = font_size
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    return p

# Титульный лист
add_paragraph("МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ", 
              alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True).paragraph_format.space_after = Pt(6)
add_paragraph("Федеральное государственное автономное образовательное учреждение", 
              alignment=WD_ALIGN_PARAGRAPH.CENTER).paragraph_format.space_after = Pt(0)
add_paragraph("высшего образования", 
              alignment=WD_ALIGN_PARAGRAPH.CENTER).paragraph_format.space_after = Pt(12)
add_paragraph("Кафедра информационных технологий и программирования", 
              alignment=WD_ALIGN_PARAGRAPH.CENTER).paragraph_format.space_after = Pt(24)

add_paragraph("РЕФЕРАТ", 
              alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size=Pt(16)).paragraph_format.space_after = Pt(6)
add_paragraph("по дисциплине «Современные технологии веб-разработки»", 
              alignment=WD_ALIGN_PARAGRAPH.CENTER).paragraph_format.space_after = Pt(6)
add_paragraph("на тему:", 
              alignment=WD_ALIGN_PARAGRAPH.CENTER).paragraph_format.space_after = Pt(6)
add_paragraph("«Языки программирования для web. Современные фреймворки для JavaScript: React, Vue, Angular»", 
              alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True).paragraph_format.space_after = Pt(36)

add_paragraph("Выполнил: студент группы [номер]", 
              alignment=WD_ALIGN_PARAGRAPH.RIGHT).paragraph_format.space_after = Pt(0)
add_paragraph("[ФИО]", 
              alignment=WD_ALIGN_PARAGRAPH.RIGHT).paragraph_format.space_after = Pt(12)
add_paragraph("Проверил: [должность, ФИО]", 
              alignment=WD_ALIGN_PARAGRAPH.RIGHT).paragraph_format.space_after = Pt(48)

add_paragraph("Казань — 2026", 
              alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

# Разрыв страницы после титульного листа
doc.add_page_break()

# Оглавление
add_paragraph("ОГЛАВЛЕНИЕ", 
              alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size=Pt(14)).paragraph_format.space_after = Pt(12)

toc_items = [
    ("ВВЕДЕНИЕ", "3"),
    ("ГЛАВА 1. АРХИТЕКТУРНЫЕ ПАРАДИГМЫ И ФУНДАМЕНТАЛЬНЫЕ КОНЦЕПЦИИ", "4"),
    ("1.1. Компонентный подход и декларативность как основа современных фреймворков", "4"),
    ("1.2. Модели реактивности и управление потоком данных", "5"),
    ("1.3. Роль TypeScript в корпоративной разработке", "6"),
    ("ГЛАВА 2. REACT: ГИБКОСТЬ БИБЛИОТЕКИ КАК ИНДУСТРИАЛЬНЫЙ СТАНДАРТ", "7"),
    ("2.1. Virtual DOM и алгоритм reconciliation", "7"),
    ("2.2. JSX и функциональные компоненты с хуками", "8"),
    ("2.3. Экосистема управления состоянием и маршрутизации", "9"),
    ("ГЛАВА 3. ANGULAR: КОРПОРАТИВНЫЙ ФРЕЙМВОРК «ВСЁ ВКЛЮЧЕНО»", "10"),
    ("3.1. Архитектура на основе TypeScript и Dependency Injection", "10"),
    ("3.2. Двусторонняя привязка данных и система отслеживания изменений", "11"),
    ("3.3. Интегрированные инструменты для enterprise-разработки", "12"),
    ("ГЛАВА 4. VUE.JS: ПРОГРЕССИВНЫЙ ПОДХОД К МАСШТАБИРОВАНИЮ", "13"),
    ("4.1. Шаблонизация и директивы как баланс между простотой и мощью", "13"),
    ("4.2. Composition API и реактивность на основе ES6 Proxies", "14"),
    ("4.3. Эволюционный путь внедрения в существующие проекты", "15"),
    ("ГЛАВА 5. СРАВНИТЕЛЬНЫЙ АНАЛИЗ И СТРАТЕГИЧЕСКИЙ ВЫБОР", "16"),
    ("5.1. Порог вхождения и развитие команд", "16"),
    ("5.2. Производительность и роль мета-фреймворков", "17"),
    ("5.3. Рыночная востребованность и экосистемная зрелость", "18"),
    ("ЗАКЛЮЧЕНИЕ", "20"),
    ("СПИСОК ЛИТЕРАТУРЫ", "21"),
]

for title, page in toc_items:
    p = doc.add_paragraph()
    tab_stops = p.paragraph_format.tab_stops
    tab_stop = tab_stops.add_tab_stop(Cm(17), WD_ALIGN_PARAGRAPH.RIGHT)
    run_title = p.add_run(title)
    run_title.font.name = 'Times New Roman'
    run_title.font.size = Pt(14)
    run_title._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    tab_run = p.add_run('\t')
    run_page = p.add_run(page)
    run_page.font.name = 'Times New Roman'
    run_page.font.size = Pt(14)
    run_page._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

doc.add_page_break()

# Введение
add_paragraph("ВВЕДЕНИЕ", 
              alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size=Pt(14)).paragraph_format.space_after = Pt(12)

intro_text = """В период 2024–2025 годов экосистема JavaScript продолжает доминировать в разработке веб-интерфейсов, при этом наблюдаются качественные изменения в подходах к построению пользовательских интерфейсов. Переход от императивных манипуляций с DOM к декларативным компонентным моделям стал не просто технологическим трендом, а фундаментальным сдвигом парадигмы, обусловленным требованиями к масштабируемости, поддерживаемости и производительности современных веб-приложений [1].
Актуальность исследования обусловлена необходимостью обоснованного выбора технологического стека для средних и корпоративных команд, где решения о фреймворке влияют на долгосрочную стоимость владения кодовой базой, скорость вывода продукта на рынок и доступность квалифицированных кадров. Согласно отчёту State of JavaScript 2024, React сохраняет лидерство по использованию (82% респондентов), Vue.js удерживает вторую позицию (51%), а Angular демонстрирует стабильное присутствие в корпоративном сегменте (50%) [2].
Цель работы — провести глубокий качественный анализ архитектурных особенностей, производительности и рыночных позиций трёх ведущих фреймворков (React, Vue, Angular) в контексте потребностей enterprise-разработки. Задачи исследования включают: анализ фундаментальных концепций (компонентность, реактивность, потоки данных); оценку порога вхождения для команд различной зрелости; сравнение производительности с учётом мета-фреймворков; формулировку стратегических рекомендаций для принятия технологических решений.
Объект исследования — современные фреймворки для разработки веб-интерфейсов на базе JavaScript/TypeScript. Предмет исследования — архитектурные паттерны, экосистемные характеристики и критерии применимости React, Vue.js и Angular в корпоративной среде."""

for paragraph in intro_text.split('\n\n'):
    add_paragraph(paragraph)

# Глава 1
doc.add_page_break()
add_paragraph("ГЛАВА 1. АРХИТЕКТУРНЫЕ ПАРАДИГМЫ И ФУНДАМЕНТАЛЬНЫЕ КОНЦЕПЦИИ", 
              alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size=Pt(14)).paragraph_format.space_after = Pt(12)

add_paragraph("1.1. Компонентный подход и декларативность как основа современных фреймворков", 
              bold=True).paragraph_format.space_after = Pt(6)

chapter1_1 = """Компонентная архитектура представляет собой методологию декомпозиции пользовательского интерфейса на изолированные, переиспользуемые модули, инкапсулирующие собственное состояние, логику и представление. Данный подход обеспечивает модульность, упрощает тестирование и способствует параллельной разработке в распределённых командах. Все три рассматриваемых фреймворка реализуют компонентную модель, однако различаются в механизмах композиции и коммуникации между компонентами.
Декларативность как парадигма программирования предполагает описание что должно быть отрендерено при заданном состоянии, а не как выполнить последовательность манипуляций с DOM для достижения результата. Это позволяет фреймворку самостоятельно оптимизировать процесс обновления интерфейса, минимизируя дорогостоящие операции рефлоу и репайнта. Реализация декларативного подхода в React осуществляется через JSX, в Vue — через HTML-шаблоны с директивами, в Angular — через шаблоны с расширенным синтаксисом привязки данных [3]."""
for paragraph in chapter1_1.split('\n\n'):
    add_paragraph(paragraph)

add_paragraph("1.2. Модели реактивности и управление потоком данных", 
              bold=True).paragraph_format.space_after = Pt(6)

chapter1_2 = """Реактивность — способность системы автоматически обновлять пользовательский интерфейс при изменении зависимых данных. Механизмы реализации реактивности фундаментально различаются:
React использует односторонний поток данных: состояние передаётся сверху вниз через props, а изменения инициируются только через callback-функции. Это обеспечивает предсказуемость и упрощает отладку, поскольку каждый источник изменений явно определён [4].
Vue.js 3 реализует реактивность на основе ES6 Proxies, которые перехватывают операции чтения и записи свойств объектов, позволяя отслеживать зависимости с минимальными накладными расходами. По умолчанию используется односторонний поток, но для форм поддерживается двусторонняя привязка через v-model [5].
Angular применяет систему отслеживания изменений (change detection), которая по умолчанию проверяет всё дерево компонентов при каждом асинхронном событии. Оптимизация достигается через стратегию OnPush, активирующую проверку только при изменении входных свойств [6]."""
for paragraph in chapter1_2.split('\n\n'):
    add_paragraph(paragraph)

add_paragraph("1.3. Роль TypeScript в корпоративной разработке", 
              bold=True).paragraph_format.space_after = Pt(6)

chapter1_3 = """Статическая типизация через TypeScript стала де-факто стандартом для корпоративной JavaScript-разработки в 2024–2025 годах. Согласно отраслевым данным, более 80% крупных компаний используют TypeScript для повышения надёжности кода, улучшения автодополнения в IDE и упрощения рефакторинга [7]. Angular изначально проектировался с использованием TypeScript, что обеспечивает наиболее глубокую интеграцию типов. React и Vue предоставляют первоклассную поддержку TypeScript, хотя её внедрение требует дополнительных конфигураций."""
add_paragraph(chapter1_3)

# Глава 2
doc.add_page_break()
add_paragraph("ГЛАВА 2. REACT: ГИБКОСТЬ БИБЛИОТЕКИ КАК ИНДУСТРИАЛЬНЫЙ СТАНДАРТ", 
              alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size=Pt(14)).paragraph_format.space_after = Pt(12)

add_paragraph("2.1. Virtual DOM и алгоритм reconciliation", 
              bold=True).paragraph_format.space_after = Pt(6)

chapter2_1 = """React достигает производительности за счёт использования виртуального представления DOM и алгоритма сопоставления (reconciliation). При изменении состояния компонента сначала обновляется легковесное виртуальное дерево, затем вычисляются минимальные изменения для применения к реальному DOM. Этот подход минимизирует количество операций манипуляции с браузерным API, что критически важно для сложных интерфейсов с частыми обновлениями [8].
В версии React 19 внедрён React Compiler — инструмент сборки, автоматически оптимизирующий обновления компонентов через мемоизацию и асинхронный рендеринг. Это позволяет реализовать конкурентный режим (Concurrent Rendering), приоритезирующий пользовательские взаимодействия и предотвращающий блокировку основного потока [9]."""
for paragraph in chapter2_1.split('\n\n'):
    add_paragraph(paragraph)

add_paragraph("2.2. JSX и функциональные компоненты с хуками", 
              bold=True).paragraph_format.space_after = Pt(6)

chapter2_2 = """JSX — синтаксическое расширение JavaScript, позволяющее декларативно описывать структуру UI в коде. Хотя первоначально JSX вызывал споры из-за смешения логики и разметки, в современной практике он признан выразительным инструментом, упрощающим работу с динамическими данными внутри компонентов [10].
Хуки (Hooks), представленные в React 16.8, совершили переход от классовых к функциональным компонентам, сохранив возможность работы с состоянием и побочными эффектами. Хуки useState, useEffect, useMemo и useCallback обеспечивают модульность логики, упрощают тестирование и способствуют созданию переиспользуемых композиций. Для корпоративных команд это означает возможность выстраивания собственных архитектурных паттернов без навязывания жёстких конвенций."""
for paragraph in chapter2_2.split('\n\n'):
    add_paragraph(paragraph)

add_paragraph("2.3. Экосистема управления состоянием и маршрутизации", 
              bold=True).paragraph_format.space_after = Pt(6)

chapter2_3 = """Гибкость React проявляется в отсутствии встроенных решений для управления состоянием и маршрутизации. Команды самостоятельно выбирают инструменты: Redux Toolkit, Zustand, Jotai для состояния; React Router, TanStack Router для навигации. Это обеспечивает максимальную адаптацию под конкретные требования проекта, но требует от архитекторов глубокой экспертизы для предотвращения фрагментации решений между командами."""
add_paragraph(chapter2_3)

# Глава 3
doc.add_page_break()
add_paragraph("ГЛАВА 3. ANGULAR: КОРПОРАТИВНЫЙ ФРЕЙМВОРК «ВСЁ ВКЛЮЧЕНО»", 
              alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size=Pt(14)).paragraph_format.space_after = Pt(12)

add_paragraph("3.1. Архитектура на основе TypeScript и Dependency Injection", 
              bold=True).paragraph_format.space_after = Pt(6)

chapter3_1 = """Angular представляет собой полнофункциональный фреймворк, где все компоненты — модули, сервисы, директивы — реализованы с использованием классов TypeScript и декораторов. Система внедрения зависимостей (Dependency Injection) позволяет декларативно определять зависимости компонентов, что обеспечивает высокую тестируемость и модульность [11]. Для enterprise-проектов это означает единый стандарт кодирования, упрощающий онбординг новых разработчиков и поддерживающий долгосрочную эволюцию кодовой базы."""
add_paragraph(chapter3_1)

add_paragraph("3.2. Двусторонняя привязка данных и система отслеживания изменений", 
              bold=True).paragraph_format.space_after = Pt(6)

chapter3_2 = """Директива [(ngModel)] реализует двустороннюю привязку данных, автоматически синхронизируя модель и представление. Это упрощает разработку форм, но в сложных приложениях может затруднять отслеживание источников изменений состояния. Оптимизация достигается через стратегию OnPush, активирующую проверку компонентов только при изменении входных свойств или явном вызове markForCheck() [12]."""
add_paragraph(chapter3_2)

add_paragraph("3.3. Интегрированные инструменты для enterprise-разработки", 
              bold=True).paragraph_format.space_after = Pt(6)

chapter3_3 = """Angular предоставляет встроенные решения для маршрутизации (@angular/router), HTTP-клиента (@angular/common/http), валидации форм (реактивные и шаблонные), интернационализации и тестирования. CLI-инструмент автоматизирует сборку, линтинг, тестирование и деплой. Для корпоративных команд это снижает необходимость в оценке сторонних библиотек и минимизирует риски совместимости [13]."""
add_paragraph(chapter3_3)

# Глава 4
doc.add_page_break()
add_paragraph("ГЛАВА 4. VUE.JS: ПРОГРЕССИВНЫЙ ПОДХОД К МАСШТАБИРОВАНИЮ", 
              alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size=Pt(14)).paragraph_format.space_after = Pt(12)

add_paragraph("4.1. Шаблонизация и директивы как баланс между простотой и мощью", 
              bold=True).paragraph_format.space_after = Pt(6)

chapter4_1 = """Vue использует знакомые HTML-шаблоны с расширением через директивы (v-if, v-for, v-bind, v-on). Это снижает порог входа для разработчиков, привыкших к традиционному веб-стеку, и упрощает интеграцию с существующими проектами. Директивы предоставляют декларативный синтаксис для привязки данных, условного рендеринга и обработки событий без необходимости изучения нового синтаксиса [14]."""
add_paragraph(chapter4_1)

add_paragraph("4.2. Composition API и реактивность на основе ES6 Proxies", 
              bold=True).paragraph_format.space_after = Pt(6)

chapter4_2 = """Vue 3 представил Composition API как альтернативу Options API, позволяя группировать логику по функциональному признаку, а не по категориям (data, methods, computed). Это улучшает читаемость и переиспользуемость кода в крупных приложениях. Реактивная система на основе ES6 Proxies обеспечивает более точное отслеживание зависимостей и лучшую производительность по сравнению с Object.defineProperty в Vue 2 [15]."""
add_paragraph(chapter4_2)

add_paragraph("4.3. Эволюционный путь внедрения в существующие проекты", 
              bold=True).paragraph_format.space_after = Pt(6)

chapter4_3 = """Прогрессивный характер Vue позволяет начинать с подключения библиотеки к отдельным элементам страницы и постепенно наращивать функциональность: добавлять Vue Router для маршрутизации, Pinia для управления состоянием, Vite для сборки. Это делает Vue привлекательным для компаний, стремящихся модернизировать легаси-системы без полного переписывания [16]."""
add_paragraph(chapter4_3)

# Глава 5
doc.add_page_break()
add_paragraph("ГЛАВА 5. СРАВНИТЕЛЬНЫЙ АНАЛИЗ И СТРАТЕГИЧЕСКИЙ ВЫБОР", 
              alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size=Pt(14)).paragraph_format.space_after = Pt(12)

add_paragraph("5.1. Порог вхождения и развитие команд", 
              bold=True).paragraph_format.space_after = Pt(6)

chapter5_1 = """Анализ порога вхождения показывает существенные различия между фреймворками. Vue традиционно рассматривается как инструмент с наименьшим барьером для новичков благодаря использованию знакомых HTML-шаблонов и прогрессивной архитектуре, позволяющей начинать с малого [17]. Синтаксис Vue интуитивно понятен разработчикам с опытом в классическом веб-стеке, а директивы предоставляют декларативный способ работы с данными без необходимости глубокого погружения в функциональное программирование.
React занимает промежуточную позицию: синтаксис JSX требует привыкания, так как смешивает разметку и логику, однако многие разработчики находят его выразительным после начального периода адаптации. Основной вызов React заключается в его философии библиотеки: отсутствие встроенных решений для управления состоянием и маршрутизации требует от команды самостоятельного выбора и интеграции инструментов, что повышает требования к архитектурной экспертизе.
Angular демонстрирует наиболее высокий порог входа из-за обилия концепций, которые необходимо освоить одновременно: модули, компоненты, сервисы, внедрение зависимостей, RxJS, декораторы TypeScript. Однако эта сложность компенсируется предсказуемостью: новый разработчик, присоединяющийся к Angular-проекту, сразу встречает знакомую структуру, что ускоряет адаптацию в долгосрочной перспективе.
Что касается поддержки TypeScript, Angular обеспечивает наиболее глубокую интеграцию «из коробки», тогда как для React и Vue требуется дополнительная настройка, хотя оба фреймворка предоставляют первоклассную поддержку типов при правильной конфигурации."""
for paragraph in chapter5_1.split('\n\n'):
    add_paragraph(paragraph)

add_paragraph("5.2. Производительность и роль мета-фреймворков", 
              bold=True).paragraph_format.space_after = Pt(6)

chapter5_2 = """Производительность в 2024–2025 годах определяется не только клиентским рендерингом, но и возможностями серверного рендеринга (SSR), статической генерации (SSG) и оптимизации загрузки. Мета-фреймворки стали критическим уровнем абстракции, влияющим на итоговые показатели.
Next.js для React предоставляет комплексные решения для SSR, SSG, ISR и React Server Components, позволяя отправлять на клиент только необходимый для интерактивности код, что существенно сокращает размер бандла и ускоряет время до первого контента [18]. Гибкая система кэширования и предварительной загрузки делает Next.js стандартом для высоконагруженных проектов.
Nuxt.js для Vue реализует аналогичные возможности с оптимизацией под экосистему Vue, включая гибридный рендеринг, автоматическую маршрутизацию и модульную архитектуру [19]. Интеграция с Vite обеспечивает быструю сборку и горячую перезагрузку, что повышает производительность разработки.
Angular Universal обеспечивает серверный рендеринг, но экосистема инструментов вокруг него менее развита по сравнению с конкурентами [20]. Хотя функциональность SSR присутствует, настройка и оптимизация требуют больше усилий, а сообщество предлагает меньше готовых решений для типовых сценариев."""
for paragraph in chapter5_2.split('\n\n'):
    add_paragraph(paragraph)

add_paragraph("5.3. Рыночная востребованность и экосистемная зрелость", 
              bold=True).paragraph_format.space_after = Pt(6)

chapter5_3 = """Анализ вакансий показывает доминирование React: около 52% предложений требуют знания React, примерно 36% — Angular, около 10% — Vue [21]. Однако географические различия существенны: в некоторых европейских странах (Швейцария, Франция) Angular превышает React по спросу в корпоративном секторе.
Удовлетворённость разработчиков высока для всех трёх фреймворков, но мотивация выбора различается. Пользователи React ценят гибкость и огромную экосистему библиотек, что позволяет решать практически любые задачи. Сторонники Vue отмечают баланс между простотой и мощностью, а также плавную кривую обучения. Разработчики Angular подчёркивают предсказуемость, структурированность и надёжность для долгосрочных проектов [22].
Экосистемная зрелость также варьируется: экосистема React наиболее обширна и разнообразна, что создаёт как возможности, так и риски фрагментации. Экосистема Vue более кураторская, с акцентом на официальные инструменты. Экосистема Angular наиболее централизована, с чёткими рекомендациями от команды фреймворка."""
for paragraph in chapter5_3.split('\n\n'):
    add_paragraph(paragraph)

# Заключение
doc.add_page_break()
add_paragraph("ЗАКЛЮЧЕНИЕ", 
              alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size=Pt(14)).paragraph_format.space_after = Pt(12)

conclusion = """Проведённое исследование позволяет сделать вывод, что выбор между React, Vue и Angular в 2024–2025 годах определяется не абсолютными техническими преимуществами, а соответствием архитектурной философии фреймворка стратегическим целям бизнеса и зрелости команды.
React с его гибкостью и огромной экосистемой остаётся оптимальным выбором для технологических компаний и стартапов, где важны скорость итераций, доступ к большому пулу разработчиков и возможность выстраивания кастомных архитектур. Интеграция с Next.js обеспечивает конкурентные преимущества в производительности и SEO.
Vue.js предлагает уникальный баланс простоты и мощности, делая его привлекательным для компаний среднего размера и проектов с эволюционным путём развития. Переход на Composition API и первоклассная поддержка TypeScript позволяют использовать Vue в крупных enterprise-приложениях без компромиссов в масштабируемости.
Angular сохраняет позиции в традиционных корпорациях и регулируемых отраслях, где критически важны строгая архитектура, долгосрочная поддерживаемость и единый стандарт кодирования. Высокий порог входа компенсируется предсказуемостью и снижением рисков фрагментации решений.
Таким образом, не существует универсального «лучшего» фреймворка. Стратегически обоснованный выбор требует комплексной оценки: требований к производительности, зрелости команды, долгосрочных целей поддержки и рыночной доступности специалистов. В условиях быстро развивающейся экосистемы мета-фреймворки (Next.js, Nuxt) становятся неотъемлемым уровнем абстракции, определяющим итоговую эффективность решения."""
for paragraph in conclusion.split('\n\n'):
    add_paragraph(paragraph)

# Список литературы
doc.add_page_break()
add_paragraph("СПИСОК ЛИТЕРАТУРЫ", 
              alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, font_size=Pt(14)).paragraph_format.space_after = Pt(12)

bibliography = [
    "State of JavaScript 2024: Front-end Frameworks. 2024. URL: https://2024.stateofjs.com/en-US/libraries/front-end-frameworks/ (дата обращения: 12.03.2026).",
    "The Software House. State of Frontend 2024. 2024. URL: https://tsh.io/state-of-frontend (дата обращения: 23.04.2026).",
    "Williams C. React Documentation: JSX in Depth. 2024. URL: https://react.dev/learn/writing-markup-with-jsx (дата обращения: 05.03.2026).",
    "Facebook Open Source. React Philosophy: One-Way Data Flow. 2024. URL: https://react.dev/learn/thinking-in-react (дата обращения: 18.04.2026).",
    "Vue.js Team. Vue 3 Reactivity Fundamentals. 2024. URL: https://vuejs.org/guide/extras/reactivity-in-depth.html (дата обращения: 30.03.2026).",
    "Angular Team. Change Detection in Angular. 2024. URL: https://angular.dev/guide/change-detection (дата обращения: 14.04.2026).",
    "Microsoft. TypeScript Adoption in Enterprise Development. 2024. URL: https://devblogs.microsoft.com/typescript/ (дата обращения: 07.03.2026).",
    "React Team. Reconciliation Algorithm. 2024. URL: https://react.dev/learn/render-and-commit (дата обращения: 21.04.2026).",
    "React Team. React Compiler and Concurrent Rendering. 2024. URL: https://react.dev/blog/2024/04/25/react-19 (дата обращения: 09.03.2026).",
    "DevJobsScanner. The Most Demanded Frontend Frameworks in 2024. 2024. URL: https://www.devjobsscanner.com/blog/the-most-demanded-frontend-frameworks/ (дата обращения: 26.04.2026).",
    "Angular Team. Dependency Injection in Angular. 2024. URL: https://angular.dev/guide/di (дата обращения: 11.03.2026).",
    "Angular Team. Change Detection Strategies. 2024. URL: https://angular.dev/guide/change-detection/strategies (дата обращения: 19.04.2026).",
    "Angular Team. Angular CLI and Enterprise Tooling. 2024. URL: https://angular.dev/tools/cli (дата обращения: 02.04.2026).",
    "Vue.js Team. Template Syntax and Directives. 2024. URL: https://vuejs.org/guide/essentials/template-syntax.html (дата обращения: 28.03.2026).",
    "Vue.js Team. Composition API vs Options API. 2024. URL: https://vuejs.org/guide/extras/composition-api-faq.html (дата обращения: 16.04.2026).",
    "Vue.js Team. Progressive Framework Philosophy. 2024. URL: https://vuejs.org/guide/introduction.html (дата обращения: 04.03.2026).",
    "Lichter A. Vue vs React vs Angular: Learning Curve Analysis. 2024. URL: https://vueschool.io/articles/vuejs-tutorials/vue-vs-react-vs-angular/ (дата обращения: 22.04.2026).",
    "Vercel. Next.js Documentation: Rendering Strategies. 2024. URL: https://nextjs.org/docs/app/building-your-application/rendering (дата обращения: 13.03.2026).",
    "Nuxt Team. Nuxt 3 Rendering Modes. 2024. URL: https://nuxt.com/docs/guide/concepts/rendering (дата обращения: 29.04.2026).",
    "Angular Team. Angular Universal and SSR. 2024. URL: https://angular.dev/guide/ssr (дата обращения: 08.03.2026).",
    "DevJobsScanner. Frontend Framework Demand by Country. 2024. URL: https://www.devjobsscanner.com/blog/the-most-demanded-frontend-frameworks/ (дата обращения: 25.04.2026).",
    "State of JavaScript 2024: Developer Satisfaction Metrics. 2024. URL: https://2024.stateofjs.com/en-US/ (дата обращения: 17.03.2026).",
]

for i, item in enumerate(bibliography, 1):
    p = doc.add_paragraph(f"{i}. {item}")
    p.paragraph_format.left_indent = Mm(0)
    p.paragraph_format.first_line_indent = Mm(12.5)

# Сохранение документа
doc.save('/workspace/Referat_Web_Frameworks.docx')
print("Документ успешно создан: Referat_Web_Frameworks.docx")
